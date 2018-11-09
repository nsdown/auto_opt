# -*- coding: utf-8 -*-
'''
@ Copyright (C) 2018        EfonFighting(email:efonfighting@126.com)(android_wechat:Efon-fighting)
@
@ description:
@
@ env stetup：
@   pip install docx
@
'''

import docx

def docxOptTemplate():
    '''
    @ 读取docx中的文本代码示例
    @
    @ return
    @
    @ param
    @ exception
    @ notice
    '''
    #获取文档对象
    file=docx.Document("/home/soy/Desktop/git.docx")
    print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

    #输出每一段的内容
    for para in file.paragraphs:
        print(para.text)
    #输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)